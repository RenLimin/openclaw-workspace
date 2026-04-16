#!/usr/bin/env python3
"""
GJSWZJ Word报告修订 v5 - XML级别样式复制
直接从模版的styles.xml复制样式定义到报告
"""
import os
import shutil
from lxml import etree
from copy import deepcopy
from docx import Document

BASE = os.path.expanduser('~/Downloads/project delivery')
TMPL_PATH = os.path.join(BASE, '（项目交付文档模板）文档名称_V1.X.docx')
REPORT_DIR = os.path.join(BASE, '20260331-第一次报告')

REPORT_FILES = [
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(Android App).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(ITS-IOS APP).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(鸿蒙NEXT APP).docx',
]

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def get_style_by_name(doc_xml, style_name):
    """Find a style element by its name attribute in styles XML."""
    for style in doc_xml:
        name_el = style.find('.//{{{}}}name'.format(W_NS))
        if name_el is not None:
            name_val = name_el.get('{{{}}}val'.format(W_NS))
            if name_val == style_name:
                return style
    return None

def copy_styles_to_report(template_doc, report_doc):
    """Copy specific styles from template to report at XML level."""
    # Target styles to copy
    target_names = [
        'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5', 'toc 6', 'toc 7', 'toc 8', 'toc 9',
        'Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5', 'Heading 6',
    ]
    
    # Get template styles XML elements
    tmpl_styles = template_doc.styles.element
    rpt_styles = report_doc.styles.element
    
    for target_name in target_names:
        tmpl_style = get_style_by_name(tmpl_styles, target_name)
        if tmpl_style is None:
            print(f"    Template style '{target_name}' not found")
            continue
        
        rpt_style = get_style_by_name(rpt_styles, target_name)
        if rpt_style is not None:
            # Replace existing style
            parent = rpt_style.getparent()
            idx = parent.index(rpt_style)
            parent.remove(rpt_style)
            new_style = deepcopy(tmpl_style)
            parent.insert(idx, new_style)
            print(f"    Replaced: {target_name}")
        else:
            # Add new style
            new_style = deepcopy(tmpl_style)
            rpt_styles.append(new_style)
            print(f"    Added: {target_name}")

def fix_paragraph_formats(report_doc):
    """Fix paragraph-level formatting to match template."""
    # Template settings:
    # Normal: font=Times New Roman, space_after=None, line_spacing=None
    # Heading 2: left_indent=266700 (1cm)
    # Heading 3: left_indent=254000
    # toc 2: left_indent=533400 (2cm)
    # toc 3: left_indent=800100 (3cm)
    # toc 2 space_after = None (was 0)
    # toc 3 space_after = None (was 0)
    
    count = 0
    for p in report_doc.paragraphs:
        changed = False
        if p.style.name == 'Normal':
            p.paragraph_format.space_after = None
            p.paragraph_format.line_spacing = None
            changed = True
        elif p.style.name == 'Heading 2':
            p.paragraph_format.left_indent = 266700
            p.paragraph_format.space_after = 76200  # template value
            changed = True
        elif p.style.name == 'Heading 3':
            p.paragraph_format.left_indent = 254000
            changed = True
        elif p.style.name == 'toc 2':
            p.paragraph_format.left_indent = 533400
            p.paragraph_format.line_spacing = None
            p.paragraph_format.space_after = None
            changed = True
        elif p.style.name == 'toc 3':
            p.paragraph_format.left_indent = 800100
            p.paragraph_format.line_spacing = None
            p.paragraph_format.space_after = None
            changed = True
        
        if changed:
            count += 1
    return count

def main():
    print("=== GJSWZJ Word报告修订 v5 (XML) ===\n")
    
    template_doc = Document(TMPL_PATH)
    print(f"Template: {TMPL_PATH}\n")
    
    for filename in REPORT_FILES:
        rpt_path = os.path.join(REPORT_DIR, filename)
        print(f"--- {filename} ---")
        
        if not os.path.exists(rpt_path):
            print(f"  SKIP: not found")
            continue
        
        # Backup
        backup = rpt_path + '.bak'
        if not os.path.exists(backup):
            shutil.copy(rpt_path, backup)
            print(f"  Backup: .bak")
        
        report_doc = Document(rpt_path)
        
        # Step 1: Copy styles at XML level
        print("  Copying styles from template...")
        copy_styles_to_report(template_doc, report_doc)
        
        # Step 2: Fix paragraph-level formatting
        print("  Fixing paragraph formatting...")
        fixed = fix_paragraph_formats(report_doc)
        print(f"    Fixed {fixed} paragraphs")
        
        # Save
        report_doc.save(rpt_path)
        print(f"  Saved\n")
    
    # Verify
    print("=== Verification ===")
    for filename in REPORT_FILES:
        rpt_path = os.path.join(REPORT_DIR, filename)
        doc = Document(rpt_path)
        print(f"\n{filename[:50]}...")
        for name in ['toc 1', 'toc 2', 'toc 3', 'Normal', 'Heading 2']:
            try:
                s = doc.styles[name]
                print(f"  {name}: size={s.font.size}, bold={s.font.bold}, "
                      f"indent={s.paragraph_format.left_indent}, "
                      f"line_spacing={s.paragraph_format.line_spacing}, "
                      f"space_after={s.paragraph_format.space_after}")
            except:
                print(f"  {name}: N/A")
    
    print("\n=== DONE ===")

if __name__ == '__main__':
    main()
