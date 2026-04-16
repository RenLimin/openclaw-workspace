#!/usr/bin/env python3
"""
GJSWZJ Word报告修订 v5 - 将模版样式应用到3个报告
"""
from docx import Document
from copy import deepcopy
from lxml import etree
import os
import shutil

BASE = os.path.expanduser('~/Downloads/project delivery')
TMPL_PATH = os.path.join(BASE, '（项目交付文档模板）文档名称_V1.X.docx')
REPORT_DIR = os.path.join(BASE, '20260331-第一次报告')

REPORT_FILES = [
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(Android App).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(ITS-IOS APP).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(鸿蒙NEXT APP).docx',
]

def copy_style_from_template(template_doc, report_doc):
    """Copy style definitions from template to report."""
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # Get template styles XML
    tmpl_styles = template_doc.part._element.xpath('//w:style')
    rpt_styles_elem = report_doc.styles.element
    
    # Build map of template styles by styleId
    tmpl_style_map = {}
    for s in tmpl_styles:
        style_id = s.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId')
        if style_id:
            tmpl_style_map[style_id] = s
    
    # Target styles to update
    target_styles = ['toc1', 'toc2', 'toc3', 'Normal', 'Heading1', 'Heading2', 'Heading3',
                     'Heading4', 'Heading5', 'Heading6', 'a0', 'a']
    
    for target_id in target_styles:
        if target_id in tmpl_style_map:
            # Find existing style in report
            existing = rpt_styles_elem.xpath(f'w:style[@w:styleId="{target_id}"]', namespaces=ns)
            if existing:
                # Replace with template style
                tmpl_style = tmpl_style_map[target_id]
                parent = existing[0].getparent()
                idx = parent.index(existing[0])
                parent.remove(existing[0])
                parent.insert(idx, deepcopy(tmpl_style))
                print(f"    Replaced style: {target_id}")
            else:
                # Add template style
                tmpl_style = tmpl_style_map[target_id]
                rpt_styles_elem.append(deepcopy(tmpl_style))
                print(f"    Added style: {target_id}")

def fix_toc_paragraphs(report_doc):
    """Ensure TOC paragraphs use correct toc styles."""
    toc_count = {'toc 1': 0, 'toc 2': 0, 'toc 3': 0}
    for p in report_doc.paragraphs:
        if p.style.name.startswith('toc'):
            toc_count[p.style.name] = toc_count.get(p.style.name, 0) + 1
    print(f"    TOC paragraphs: {toc_count}")

def fix_heading_indents(report_doc):
    """Fix Heading 2 indent to match template (266700 = 1cm)."""
    # Template Heading 2: left_indent = 266700
    # Template Heading 3: left_indent = 254000
    count = 0
    for p in report_doc.paragraphs:
        if p.style.name == 'Heading 2':
            p.paragraph_format.left_indent = 266700
            count += 1
        elif p.style.name == 'Heading 3':
            p.paragraph_format.left_indent = 254000
            count += 1
    print(f"    Fixed {count} heading indents")

def main():
    print("=== GJSWZJ Word报告修订 v5 ===\n")
    
    template_doc = Document(TMPL_PATH)
    print(f"Template loaded: {TMPL_PATH}\n")
    
    for i, filename in enumerate(REPORT_FILES):
        rpt_path = os.path.join(REPORT_DIR, filename)
        print(f"--- Processing: {filename} ---")
        
        if not os.path.exists(rpt_path):
            print(f"  SKIP: File not found")
            continue
        
        # Backup original
        backup_path = rpt_path + '.bak'
        if not os.path.exists(backup_path):
            shutil.copy(rpt_path, backup_path)
            print(f"  Backup created: .bak")
        
        # Load report
        report_doc = Document(rpt_path)
        
        # Step 1: Copy template styles to report
        print("  Applying template styles...")
        copy_style_from_template(template_doc, report_doc)
        
        # Step 2: Fix TOC paragraph formatting
        print("  Checking TOC paragraphs...")
        fix_toc_paragraphs(report_doc)
        
        # Step 3: Fix heading indents
        print("  Fixing heading indents...")
        fix_heading_indents(report_doc)
        
        # Step 4: Fix Normal style
        print("  Fixing Normal style...")
        try:
            normal = report_doc.styles['Normal']
            normal.font.name = 'Times New Roman'
            normal.paragraph_format.space_after = None
            normal.paragraph_format.line_spacing = None
        except Exception as e:
            print(f"    Error fixing Normal: {e}")
        
        # Save
        report_doc.save(rpt_path)
        print(f"  Saved: {filename}")
        print()
    
    print("=== DONE ===")

if __name__ == '__main__':
    main()
