#!/usr/bin/env python3
"""
能力 4: Word 解析测试
测试 python-docx: TOC、标题级别、段落、表格
"""
import os
import sys

def test_python_docx_basic():
    """测试 4.1: python-docx 基础解析"""
    print("📝 测试 Word 基础解析...")
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Create test document
        doc = Document()
        doc.add_heading('Test Document', level=1)
        doc.add_paragraph('This is the first paragraph.')
        doc.add_paragraph('This is the second paragraph.')
        
        test_file = "/tmp/test-word.docx"
        doc.save(test_file)
        
        # Re-open and parse
        doc2 = Document(test_file)
        print(f"  段落数: {len(doc2.paragraphs)}")
        print(f"  第一段: {doc2.paragraphs[0].text}")
        assert len(doc2.paragraphs) == 3
        assert doc2.paragraphs[0].text == "Test Document"
        print("  ✅ Word 基础解析通过")
        return True
    except Exception as e:
        print(f"  ❌ Word 基础解析失败: {e}")
        return False

def test_heading_levels():
    """测试 4.2: 标题级别识别"""
    print("\n📑 测试标题级别识别...")
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Create document with multiple heading levels
        doc = Document()
        doc.add_heading('主标题', level=1)
        doc.add_paragraph('引言段落')
        doc.add_heading('第一章', level=2)
        doc.add_paragraph('第一章内容')
        doc.add_heading('1.1 节', level=3)
        doc.add_paragraph('1.1 节内容')
        doc.add_heading('第二章', level=2)
        doc.add_paragraph('第二章内容')
        
        test_file = "/tmp/test-word-headings.docx"
        doc.save(test_file)
        
        # Re-open and parse
        doc2 = Document(test_file)
        headings = []
        for para in doc2.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 0
                headings.append({"text": para.text, "level": level})
        
        print(f"  找到 {len(headings)} 个标题:")
        for h in headings:
            print(f"    H{h['level']}: {h['text']}")
        assert len(headings) == 4
        assert headings[0]["level"] == 1
        assert headings[1]["level"] == 2
        assert headings[2]["level"] == 3
        print("  ✅ 标题级别识别通过")
        return True
    except Exception as e:
        print(f"  ❌ 标题级别识别失败: {e}")
        return False

def test_toc_extraction():
    """测试 4.3: TOC（目录）提取"""
    print("\n📋 测试 TOC 提取...")
    try:
        from docx import Document
        
        doc = Document("/tmp/test-word-headings.docx")
        toc = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 0
                indent = "  " * (level - 1)
                toc.append(f"{indent}{para.text}")
        
        print("  目录:")
        for item in toc:
            print(f"    {item}")
        assert len(toc) == 4
        print("  ✅ TOC 提取通过")
        return True
    except Exception as e:
        print(f"  ❌ TOC 提取失败: {e}")
        return False

def test_tables():
    """测试 4.4: 表格提取"""
    print("\n📊 测试表格提取...")
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Create document with table
        doc = Document()
        doc.add_heading('包含表格的文档', level=1)
        table = doc.add_table(rows=3, cols=3, style='Table Grid')
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(0, 2).text = "City"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        table.cell(1, 2).text = "Beijing"
        table.cell(2, 0).text = "Bob"
        table.cell(2, 1).text = "25"
        table.cell(2, 2).text = "Shanghai"
        
        test_file = "/tmp/test-word-table.docx"
        doc.save(test_file)
        
        # Re-open and parse
        doc2 = Document(test_file)
        tables = doc2.tables
        print(f"  表格数量: {len(tables)}")
        if tables:
            table = tables[0]
            print(f"  表格尺寸: {len(table.rows)}x{len(table.columns)}")
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                print(f"    {cells}")
        assert len(tables) == 1
        assert len(tables[0].rows) == 3
        print("  ✅ 表格提取通过")
        return True
    except Exception as e:
        print(f"  ❌ 表格提取失败: {e}")
        return False

def main():
    print("=" * 50)
    print("能力 4: Word 解析测试")
    print("=" * 50)
    
    results = {}
    results["test_4_1_basic"] = "✅ 通过" if test_python_docx_basic() else "❌ 失败"
    results["test_4_2_headings"] = "✅ 通过" if test_heading_levels() else "❌ 失败"
    results["test_4_3_toc"] = "✅ 通过" if test_toc_extraction() else "❌ 失败"
    results["test_4_4_tables"] = "✅ 通过" if test_tables() else "❌ 失败"
    
    print("\n" + "=" * 50)
    print("测试结果汇总:")
    for k, v in results.items():
        print(f"  {k}: {v}")
    passed = sum(1 for v in results.values() if "✅" in v)
    print(f"  通过率: {passed}/{len(results)}")
    print("=" * 50)
    
    return passed == len(results)

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
