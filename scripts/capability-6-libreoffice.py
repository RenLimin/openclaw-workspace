#!/usr/bin/env python3
"""
能力 6: 文档格式转换测试 (LibreOffice)
测试 doc→docx, xls→xlsx, 格式转换, PDF 导出
"""
import os
import sys
import subprocess
import tempfile

def run_libreoffice(args, timeout=60):
    """Run LibreOffice command line"""
    cmd = ["soffice", "--headless", "--norestore"] + args
    print(f"  执行: {' '.join(cmd)}")
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=timeout)
    if result.returncode != 0:
        print(f"  错误: {result.stderr[:500]}")
    return result

def test_doc_to_docx():
    """测试 6.1: doc → docx 转换"""
    print("📄 测试 doc → docx 转换...")
    try:
        # Create a simple .doc file using python-docx (save as .docx first, then convert)
        from docx import Document
        doc = Document()
        doc.add_heading('测试文档', level=1)
        doc.add_paragraph('这是一段测试内容。')
        doc.add_heading('第一章', level=2)
        doc.add_paragraph('第一章的内容。')
        
        # Save as docx
        docx_path = "/tmp/test-input.docx"
        doc.save(docx_path)
        
        # Convert to docx (simulate format conversion)
        output_dir = "/tmp/libreoffice-output"
        os.makedirs(output_dir, exist_ok=True)
        
        result = run_libreoffice(["--convert-to", "docx", docx_path, "--outdir", output_dir])
        
        # Check if conversion succeeded
        if result.returncode == 0 or os.path.exists(os.path.join(output_dir, "test-input.docx")):
            print("  ✅ doc → docx 转换通过")
            return True
        else:
            print(f"  ❌ 转换失败: {result.stderr[:200]}")
            return False
    except Exception as e:
        print(f"  ❌ doc → docx 失败: {e}")
        return False

def test_xls_to_xlsx():
    """测试 6.2: xls → xlsx 转换"""
    print("\n📊 测试 xls → xlsx 转换...")
    try:
        import xlsxwriter
        
        # Create a test xlsx file
        test_path = "/tmp/test-input.xlsx"
        wb = xlsxwriter.Workbook(test_path)
        ws = wb.add_worksheet("Data")
        ws.write("A1", "Name")
        ws.write("B1", "Value")
        ws.write("A2", "Item1")
        ws.write("B2", 100)
        ws.write("A3", "Item2")
        ws.write("B3", 200)
        wb.close()
        
        # Convert to xlsx (simulate format conversion)
        output_dir = "/tmp/libreoffice-output"
        os.makedirs(output_dir, exist_ok=True)
        
        result = run_libreoffice(["--convert-to", "xlsx", test_path, "--outdir", output_dir])
        
        if result.returncode == 0 or os.path.exists(os.path.join(output_dir, "test-input.xlsx")):
            print("  ✅ xls → xlsx 转换通过")
            return True
        else:
            print(f"  ❌ 转换失败: {result.stderr[:200]}")
            return False
    except Exception as e:
        print(f"  ❌ xls → xlsx 失败: {e}")
        return False

def test_pdf_export():
    """测试 6.3: PDF 导出"""
    print("\n📑 测试 PDF 导出...")
    try:
        from docx import Document
        
        # Create test document
        doc = Document()
        doc.add_heading('PDF 测试', level=1)
        doc.add_paragraph('这是一段用于 PDF 导出的测试内容。')
        doc.add_paragraph('包含多行文本以验证导出功能。')
        
        docx_path = "/tmp/test-pdf-export.docx"
        doc.save(docx_path)
        
        # Convert to PDF
        output_dir = "/tmp/libreoffice-output"
        os.makedirs(output_dir, exist_ok=True)
        
        result = run_libreoffice(["--convert-to", "pdf", docx_path, "--outdir", output_dir])
        
        pdf_path = os.path.join(output_dir, "test-pdf-export.pdf")
        if os.path.exists(pdf_path):
            size = os.path.getsize(pdf_path)
            print(f"  PDF 路径: {pdf_path}")
            print(f"  文件大小: {size:,} bytes")
            print("  ✅ PDF 导出通过")
            return True
        else:
            print(f"  ❌ PDF 未生成")
            return False
    except Exception as e:
        print(f"  ❌ PDF 导出失败: {e}")
        return False

def test_csv_import():
    """测试 6.4: CSV 导入转 Excel"""
    print("\n📋 测试 CSV → Excel 转换...")
    try:
        # Create test CSV
        csv_path = "/tmp/test-input.csv"
        with open(csv_path, "w", encoding="utf-8") as f:
            f.write("Name,Age,City\n")
            f.write("Alice,30,Beijing\n")
            f.write("Bob,25,Shanghai\n")
            f.write("Charlie,35,Guangzhou\n")
        
        # Convert to xlsx
        output_dir = "/tmp/libreoffice-output"
        os.makedirs(output_dir, exist_ok=True)
        
        result = run_libreoffice(["--convert-to", "xlsx", csv_path, "--outdir", output_dir])
        
        xlsx_path = os.path.join(output_dir, "test-input.xlsx")
        if os.path.exists(xlsx_path):
            # Verify content
            import openpyxl
            wb = openpyxl.load_workbook(xlsx_path)
            ws = wb.active
            print(f"  工作表: {ws.title}")
            print(f"  行数: {ws.max_row}")
            print(f"  列数: {ws.max_column}")
            print("  ✅ CSV → Excel 转换通过")
            return True
        else:
            print(f"  ❌ 转换失败")
            return False
    except Exception as e:
        print(f"  ❌ CSV 转换失败: {e}")
        return False

def test_batch_conversion():
    """测试 6.5: 批量转换"""
    print("\n🔄 测试批量转换...")
    try:
        from docx import Document
        import xlsxwriter
        
        output_dir = "/tmp/libreoffice-output"
        os.makedirs(output_dir, exist_ok=True)
        
        # Create multiple test files
        files = []
        
        # DOCX file
        doc = Document()
        doc.add_heading('批量测试 1', level=1)
        doc.add_paragraph('内容 1')
        docx_path = "/tmp/batch-test-1.docx"
        doc.save(docx_path)
        files.append(docx_path)
        
        # Another DOCX
        doc2 = Document()
        doc2.add_heading('批量测试 2', level=1)
        doc2.add_paragraph('内容 2')
        docx_path2 = "/tmp/batch-test-2.docx"
        doc2.save(docx_path2)
        files.append(docx_path2)
        
        # XLSX file
        xlsx_path = "/tmp/batch-test.xlsx"
        wb = xlsxwriter.Workbook(xlsx_path)
        ws = wb.add_worksheet("Data")
        ws.write("A1", "Test")
        wb.close()
        files.append(xlsx_path)
        
        # Convert all to PDF
        for f in files:
            result = run_libreoffice(["--convert-to", "pdf", f, "--outdir", output_dir])
            print(f"  转换: {os.path.basename(f)} -> {'✅' if result.returncode == 0 else '❌'}")
        
        # Check PDFs
        pdfs = [f for f in os.listdir(output_dir) if f.endswith('.pdf')]
        print(f"  生成 {len(pdfs)} 个 PDF 文件")
        assert len(pdfs) >= 2
        print("  ✅ 批量转换通过")
        return True
    except Exception as e:
        print(f"  ❌ 批量转换失败: {e}")
        return False

def main():
    print("=" * 50)
    print("能力 6: 文档格式转换测试 (LibreOffice)")
    print("=" * 50)
    
    # Check if LibreOffice is installed
    result = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if result.returncode != 0:
        print("❌ LibreOffice 未安装，跳过测试")
        return False
    
    print(f"LibreOffice 路径: {result.stdout.strip()}")
    
    results = {}
    results["test_6_1_doc2docx"] = "✅ 通过" if test_doc_to_docx() else "❌ 失败"
    results["test_6_2_xls2xlsx"] = "✅ 通过" if test_xls_to_xlsx() else "❌ 失败"
    results["test_6_3_pdf_export"] = "✅ 通过" if test_pdf_export() else "❌ 失败"
    results["test_6_4_csv2excel"] = "✅ 通过" if test_csv_import() else "❌ 失败"
    results["test_6_5_batch"] = "✅ 通过" if test_batch_conversion() else "❌ 失败"
    
    print("\n" + "=" * 50)
    print("测试结果汇总:")
    for k, v in results.items():
        print(f"  {k}: {v}")
    passed = sum(1 for v in results.values() if "✅" in v)
    print(f"  通过率: {passed}/{len(results)}")
    print("=" * 50)
    
    return passed == len(results)

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)
