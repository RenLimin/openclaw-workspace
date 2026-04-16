#!/usr/bin/env python3
"""
GJSWZJ 项目交付报告修订脚本
修复字体、段落间距、TOC 一致性
"""

import copy
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

TEMPLATE_PATH = os.path.expanduser('~/Downloads/project delivery/（项目交付文档模板）文档名称_V1.X.docx')
REPORT_DIR = os.path.expanduser('~/Downloads/project delivery/20260331-第一次报告/')
OUTPUT_DIR = os.path.expanduser('~/Downloads/openclaw-skill/')

REPORTS = [
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(Android App).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(ITS-IOS APP).docx',
    '国家税务总局2026年自然人税收管理系统升级完善及运行维护项目_APP测评报告第1次(鸿蒙NEXT APP).docx',
]

def copy_style_elements(src_doc, dst_doc, style_name):
    """Copy style element from source to destination document"""
    src_style = src_doc.styles[style_name] if style_name in src_doc.styles else None
    dst_style = dst_doc.styles[style_name] if style_name in dst_doc.styles else None
    if src_style is None or dst_style is None:
        return
    
    # Copy paragraph properties
    src_pPr = src_style.element.find(qn('w:pPr'))
    dst_pPr = dst_style.element.find(qn('w:pPr'))
    
    if src_pPr is not None:
        if dst_pPr is not None:
            # Replace pPr content
            for child in list(dst_pPr):
                dst_pPr.remove(child)
            for child in src_pPr:
                dst_pPr.append(copy.deepcopy(child))
        else:
            dst_style.element.append(copy.deepcopy(src_pPr))
    
    # Copy run properties
    src_rPr = src_style.element.find(qn('w:rPr'))
    dst_rPr = dst_style.element.find(qn('w:rPr'))
    
    if src_rPr is not None:
        if dst_rPr is not None:
            for child in list(dst_rPr):
                dst_rPr.remove(child)
            for child in src_rPr:
                dst_rPr.append(copy.deepcopy(child))
        else:
            # Create rPr
            if dst_pPr is None:
                dst_pPr = dst_style.element.makeelement(qn('w:pPr'), {})
                dst_style.element.append(dst_pPr)
            idx = list(dst_style.element).index(dst_pPr)
            rPr = dst_style.element.makeelement(qn('w:rPr'), {})
            dst_style.element.insert(idx, rPr)
            for child in src_rPr:
                rPr.append(copy.deepcopy(child))

def normalize_body_text(doc):
    """Normalize body text: font to 宋体, spacing to 1.5 line"""
    count = 0
    for p in doc.paragraphs:
        if p.style.name not in ('Normal', '正文应用', 'Caption'):
            continue
        if p.style.name == 'Caption':
            continue
        # Skip TOC entries
        if p.style.name.startswith('toc'):
            continue
        # Skip empty paragraphs used for spacing
        if not p.text.strip():
            continue
        
        # Normalize font to 宋体
        for run in p.runs:
            if run.font.name is not None and run.font.name not in ('宋体', 'Times New Roman'):
                run.font.name = '宋体'
            if run.font.size is not None and run.font.size > 200000:
                # Keep heading-sized text as-is (these are cover elements)
                pass
            elif run.font.size is None:
                run.font.size = Pt(10.5)  # 10.5pt default for body text
        
        # Normalize spacing to 1.5 line
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        
        count += 1
    return count

def normalize_table_text(doc):
    """Normalize table cell text to 宋体, 10.5pt"""
    count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.name is None or run.font.name == '方正小标宋_GBK':
                            run.font.name = '宋体'
                        if run.font.size is None:
                            run.font.size = Pt(10.5)
                    count += 1
    return count

def normalize_heading_spacing(doc, template_doc):
    """Normalize heading spacing to match template"""
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        copy_style_elements(template_doc, doc, style_name)

def update_toc(doc):
    """Ensure TOC field is present and properly formatted"""
    # Find existing TOC paragraphs and ensure they match template style
    for p in doc.paragraphs:
        if p.style.name.startswith('toc'):
            # TOC entries should use 宋体 font
            for run in p.runs:
                if run.font.name is None:
                    run.font.name = '宋体'

def main():
    template_doc = Document(TEMPLATE_PATH)
    
    for report_name in REPORTS:
        src_path = os.path.join(REPORT_DIR, report_name)
        output_path = os.path.join(OUTPUT_DIR, report_name)
        
        print(f"\n=== Processing: {report_name[:60]}... ===")
        doc = Document(src_path)
        
        # 1. Copy heading styles from template
        normalize_heading_spacing(doc, template_doc)
        
        # 2. Normalize body text font and spacing
        body_count = normalize_body_text(doc)
        print(f"  Normalized {body_count} body text paragraphs")
        
        # 3. Normalize table text
        table_count = normalize_table_text(doc)
        print(f"  Normalized {table_count} table cell paragraphs")
        
        # 4. Update TOC styling
        update_toc(doc)
        
        # 5. Save
        doc.save(output_path)
        print(f"  Saved to: {output_path}")
    
    print("\nDone!")

if __name__ == '__main__':
    main()
